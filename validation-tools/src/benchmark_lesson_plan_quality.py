#!/usr/bin/env python3

from __future__ import annotations

import copy
import json
import os
import re
import subprocess
import tempfile
from pathlib import Path
from typing import Any

from docx import Document


WORKSPACE_ROOT = Path(__file__).resolve().parents[2]
LAUNCHER = WORKSPACE_ROOT / "Generate Lesson Plan.command"
SAMPLE_DECK = WORKSPACE_ROOT / "codex-lesson-plan-generator" / "examples" / "sample_input_slides.pptx"
GOLD_STANDARD_DIR = WORKSPACE_ROOT / "lesson-plan-engine" / "gold-standards" / "sample-triangle-launch"
DATE_PATTERN = re.compile(r"\b\d{4}-\d{2}-\d{2}\b")
LONG_DATE_PATTERN = re.compile(
    r"\b(?:January|February|March|April|May|June|July|August|September|October|November|December) "
    r"\d{1,2}, \d{4}\b"
)


def load_json(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def load_docx_paragraphs(path: Path) -> list[str]:
    document = Document(str(path))
    return [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]


def load_docx_body_blocks(path: Path) -> list[dict[str, Any]]:
    document = Document(str(path))
    blocks: list[dict[str, Any]] = []
    for child in document.element.body.iterchildren():
        tag = child.tag.rsplit("}", 1)[-1]
        if tag == "p":
            text_parts = [
                node.text
                for node in child.iter()
                if node.tag.rsplit("}", 1)[-1] == "t" and node.text
            ]
            text = normalize_text(" ".join(" ".join(text_parts).split()))
            if text:
                blocks.append({"type": "paragraph", "text": text})
        elif tag == "tbl":
            table_rows = []
            for row in child.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tr"):
                cells = row.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tc")
                row_values = []
                for cell in cells:
                    text_parts = [
                        node.text
                        for node in cell.iter()
                        if node.tag.rsplit("}", 1)[-1] == "t" and node.text
                    ]
                    row_values.append(normalize_text(" ".join(" ".join(text_parts).split())))
                if row_values:
                    table_rows.append(row_values)
            if table_rows:
                blocks.append({"type": "table", "rows": table_rows})
    return blocks


def normalize_lesson_plan_json(data: dict[str, Any]) -> dict[str, Any]:
    normalized = copy.deepcopy(data)
    normalized["date"] = "<DATE>"
    source_deck = normalized.get("source_deck", {})
    if isinstance(source_deck, dict):
        source_deck["source_file"] = Path(str(source_deck.get("source_file", ""))).name
    for session in normalized.get("sessions", []):
        lesson_information = session.get("lesson_information", {})
        if isinstance(lesson_information, dict) and "date" in lesson_information:
            lesson_information["date"] = "<DATE>"
    return normalized


def normalize_text(text: str) -> str:
    normalized = DATE_PATTERN.sub("<DATE>", text)
    return LONG_DATE_PATTERN.sub("<DATE>", normalized)


def normalize_docx_paragraphs(paragraphs: list[str]) -> list[str]:
    return [normalize_text(paragraph) for paragraph in paragraphs]


def require_path(path: Path) -> None:
    if not path.exists():
        raise SystemExit(f"Missing required lesson-plan benchmark path: {path}")


def compare_outputs(output_dir: Path) -> list[str]:
    failures: list[str] = []

    expected_json = normalize_lesson_plan_json(load_json(GOLD_STANDARD_DIR / "lesson_plan.json"))
    actual_json = normalize_lesson_plan_json(load_json(output_dir / "lesson_plan.json"))
    if actual_json != expected_json:
        failures.append("lesson_plan.json did not match the promoted gold standard after date/path normalization")

    expected_md = normalize_text((GOLD_STANDARD_DIR / "lesson_plan.md").read_text(encoding="utf-8"))
    actual_md = normalize_text((output_dir / "lesson_plan.md").read_text(encoding="utf-8"))
    if actual_md != expected_md:
        failures.append("lesson_plan.md did not match the promoted gold standard after date normalization")

    expected_report = (GOLD_STANDARD_DIR / "validation_report.md").read_text(encoding="utf-8")
    actual_report = (output_dir / "validation_report.md").read_text(encoding="utf-8")
    if actual_report != expected_report:
        failures.append("validation_report.md did not match the promoted gold standard")

    expected_docx = normalize_docx_paragraphs(load_docx_paragraphs(GOLD_STANDARD_DIR / "lesson_plan.docx"))
    actual_docx = normalize_docx_paragraphs(load_docx_paragraphs(output_dir / "lesson_plan.docx"))
    if actual_docx != expected_docx:
        failures.append("lesson_plan.docx text did not match the promoted gold standard after date normalization")

    expected_docx_blocks = load_docx_body_blocks(GOLD_STANDARD_DIR / "lesson_plan.docx")
    actual_docx_blocks = load_docx_body_blocks(output_dir / "lesson_plan.docx")
    if actual_docx_blocks != expected_docx_blocks:
        failures.append("lesson_plan.docx body structure did not match the promoted gold standard after date normalization")

    return failures


def main() -> int:
    for required_path in (
        LAUNCHER,
        SAMPLE_DECK,
        GOLD_STANDARD_DIR / "lesson_plan.json",
        GOLD_STANDARD_DIR / "lesson_plan.md",
        GOLD_STANDARD_DIR / "validation_report.md",
        GOLD_STANDARD_DIR / "lesson_plan.docx",
    ):
        require_path(required_path)

    with tempfile.TemporaryDirectory(prefix="lesson_plan_quality_benchmark_") as temp_dir:
        root = Path(temp_dir)
        inbox_dir = root / "INBOX"
        output_dir = root / "OUTPUT"
        inbox_dir.mkdir(parents=True, exist_ok=True)
        output_dir.mkdir(parents=True, exist_ok=True)

        env = os.environ.copy()
        env["LESSON_PLAN_NO_OPEN"] = "1"
        env["EDUWONDERLAB_WATCH_DIR"] = str(inbox_dir)
        env["EDUWONDERLAB_OUTPUT_DIR"] = str(output_dir)

        subprocess.run(
            [str(LAUNCHER), "--deck", str(SAMPLE_DECK)],
            cwd=str(WORKSPACE_ROOT),
            env=env,
            check=True,
        )

        failures = compare_outputs(output_dir)
        if failures:
            raise SystemExit("Lesson plan quality benchmark failed:\n- " + "\n- ".join(failures))

    print("PASS lesson_plan_sample_triangle_launch")
    print("Lesson plan quality benchmark passed.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
