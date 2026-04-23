from __future__ import annotations

from datetime import datetime
import re
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from apply_supports import student_requires_small_group
from utils import clean_line, ensure_directory, select_preferred_learning_target, unique_preserve, write_text


DEFAULT_VOCAB_DEFINITIONS = {
    "area": "The number of square units that cover a figure.",
    "triangle": "A polygon with three sides and three angles.",
    "rectangle": "A quadrilateral with four right angles.",
    "parallelogram": "A quadrilateral with two pairs of parallel sides.",
    "trapezoid": "A quadrilateral with at least one pair of parallel sides.",
    "rhombus": "A parallelogram with four congruent sides.",
    "octagon": "A polygon with eight sides.",
    "polygon": "A closed figure made of straight line segments.",
    "regular polygon": "A polygon with equal side lengths and equal angle measures.",
    "composite figure": "A figure made from two or more simpler shapes.",
    "diagonal": "A line segment connecting two nonconsecutive vertices of a polygon.",
    "decompose": "To break a figure into smaller, familiar shapes.",
    "compose": "To combine smaller shapes into a new figure.",
    "base": "The side used with the height to measure area.",
    "height": "The perpendicular distance from the base to the opposite vertex or side.",
    "equation": "A mathematical statement showing two expressions are equal.",
    "inequality": "A statement that compares quantities using symbols such as >, <, >=, or <=.",
    "variable": "A symbol that represents an unknown quantity.",
    "independent variable": "The input value that you choose or control first.",
    "dependent variable": "The output value that changes because of the input.",
    "constant of proportionality": "The number multiplied by the input in a proportional relationship.",
    "proportional relationship": "A relationship in which the output equals a constant rate times the input.",
    "rate": "The amount that changes for each one unit of input.",
    "input": "The value you start with before the rule is applied.",
    "output": "The value produced after the rule is applied.",
}

SPANISH_TRANSLATIONS = {
    "area": "area",
    "triangle": "triangulo",
    "rectangle": "rectangulo",
    "parallelogram": "paralelogramo",
    "trapezoid": "trapecio",
    "rhombus": "rombo",
    "octagon": "octagono",
    "polygon": "poligono",
    "regular polygon": "poligono regular",
    "composite figure": "figura compuesta",
    "diagonal": "diagonal",
    "decompose": "descomponer",
    "compose": "componer",
    "base": "base",
    "height": "altura",
    "equation": "ecuacion",
    "inequality": "desigualdad",
    "variable": "variable",
    "independent variable": "variable independiente",
    "dependent variable": "variable dependiente",
    "constant of proportionality": "constante de proporcionalidad",
    "proportional relationship": "relacion proporcional",
    "rate": "tasa",
    "input": "entrada",
    "output": "salida",
}

MORPHOLOGY_HINTS = {
    "area": "root area",
    "triangle": "tri- + angle",
    "rectangle": "rect- + angle",
    "parallelogram": "parallel + -gram",
    "trapezoid": "trapez- + -oid",
    "rhombus": "rhomb- + -us",
    "octagon": "octa- + -gon",
    "polygon": "poly- + -gon",
    "regular polygon": "regular + poly- + -gon",
    "composite figure": "com- + posit + figure",
    "diagonal": "dia- + gon",
    "decompose": "de- + compose",
    "compose": "com- + pose",
    "base": "base/root",
    "height": "height/root",
    "equation": "equa- + -tion",
    "inequality": "in- + equal + -ity",
    "variable": "vary + -able",
    "independent variable": "in- + depend",
    "dependent variable": "de- + pend",
    "constant of proportionality": "constant + proportion",
    "proportional relationship": "pro- + portion",
    "rate": "rate/root",
    "input": "in + put",
    "output": "out + put",
}

CROSS_DISCIPLINARY_HINTS = {
    "area": "Geometry, Science",
    "triangle": "Geometry, Engineering",
    "rectangle": "Geometry, Design",
    "parallelogram": "Geometry, Design",
    "trapezoid": "Geometry, Architecture",
    "rhombus": "Geometry, Design",
    "octagon": "Geometry, Architecture",
    "polygon": "Geometry, Computer Graphics",
    "regular polygon": "Geometry, Art",
    "composite figure": "Geometry, Design",
    "diagonal": "Geometry, Engineering",
    "decompose": "Geometry, Problem Solving",
    "compose": "Geometry, Design",
    "equation": "Algebra, Science formulas",
    "inequality": "Algebra, Economics",
    "variable": "Algebra, Science",
    "independent variable": "Science (manipulated variable)",
    "dependent variable": "Science (responding variable)",
    "constant of proportionality": "Algebra, Economics",
    "proportional relationship": "Algebra, Science",
    "rate": "Science, Economics",
    "input": "Computer Science, Functions",
    "output": "Computer Science, Functions",
}

PROCEDURE_ROW_ORDER = [
    ("warm_up", "Warm-Up"),
    ("launch", "Launch"),
    ("be_curious", "Be Curious"),
    ("guided_practice", "Guided Practice"),
    ("collaborative_practice", "Collaborative Practice"),
    ("independent_practice", "Independent Practice"),
    ("closing", "Closing"),
    ("exit_ticket", "Exit Ticket"),
]

SUPPORT_PRIORITIES = {
    "warm_up": ["sentence starters", "word banks", "clarified/repeated directions", "visuals"],
    "launch": ["sentence starters", "word banks", "visuals", "graphic organizers"],
    "be_curious": ["sentence starters", "word banks", "expressive language support", "visuals"],
    "guided_practice": ["graphic organizers", "visuals", "immediate feedback", "calculator"],
    "collaborative_practice": ["small group", "sentence starters", "word banks", "reduced distractions"],
    "independent_practice": ["small group", "extended time", "chunking", "calculator", "text-to-speech"],
    "closing": ["clarified/repeated directions", "sentence starters", "word banks", "visuals"],
    "exit_ticket": ["extended time", "word banks", "sentence starters", "calculator"],
}

GRADE_ORDINALS = {1: "1st", 2: "2nd", 3: "3rd"}

STANDARD_ALIGNMENT_MAP = {
    "6.EE.C.9": {
        "prior": "5.OA.B.3",
        "next": "7.RP.A.2",
    },
    "6.G.A.1": {
        "prior": "5.MD.C.5",
        "next": "7.G.B.6",
    },
}

STANDARD_DESCRIPTION_OVERRIDES = {
    "5.MD.C.5": (
        "Relate volume to the operations of multiplication and addition and solve real-world and "
        "mathematical problems involving volume."
    ),
    "5.OA.B.3": "Generate two numerical patterns using given rules. Identify relationships between corresponding terms.",
    "6.EE.C.9": (
        "Use variables to represent two quantities in a real-world problem that change in relationship "
        "to one another; write an equation to express one quantity, thought of as the dependent variable, "
        "in terms of the other quantity, thought of as the independent variable; analyze the relationship "
        "using graphs and tables, and relate these to the equation."
    ),
    "6.G.A.1": (
        "Find the area of right triangles, other triangles, special quadrilaterals, and polygons by "
        "composing into rectangles or decomposing into triangles and other shapes; apply these techniques "
        "in the context of solving real-world and mathematical problems."
    ),
    "7.G.B.6": (
        "Solve real-world and mathematical problems involving area, volume and surface area of two- and "
        "three-dimensional objects composed of triangles, quadrilaterals, polygons, cubes, and right prisms."
    ),
    "7.RP.A.2": "Recognize and represent proportional relationships between quantities.",
}

GENERIC_PHASE_LINES = {
    "be curious",
    "closure",
    "discussion",
    "do now",
    "exit ticket",
    "guided practice",
    "independent practice",
    "launch",
    "learning target",
    "lesson launch",
    "mini-lesson",
    "model",
    "session 1",
    "session 2",
    "warm-up",
    "warm up",
    "workspace",
}


def ensure_template_docx(template_path: Path, config: dict[str, Any]) -> None:
    if template_path.exists():
        return

    ensure_directory(template_path.parent)
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    _apply_base_styles(document, config)
    document.save(str(template_path))


def render_markdown(
    lesson_plan: dict[str, Any],
    template_path: Path,
    output_path: Path,
    config: dict[str, Any],
) -> None:
    del template_path
    lines: list[str] = []
    for index, session in enumerate(lesson_plan.get("sessions", [])):
        if index:
            lines.extend(["", "---", ""])
        lines.extend(_render_session_markdown(session, config))
    write_text(output_path, "\n".join(lines).strip() + "\n")


def render_docx(
    lesson_plan: dict[str, Any],
    template_path: Path,
    output_path: Path,
    config: dict[str, Any],
) -> list[Path]:
    output_dir = output_path.parent
    ensure_directory(output_dir)
    rendered_paths: list[Path] = []

    # --- CLOUD TRIGGER: Create Google Doc ---
    gas_url = "https://script.google.com/macros/s/AKfycbwaylDmvzetnuWLxbT5IFE0GXVSIyUuWSPFK0EadYiRoT_4uDTTU1nKYd3TRZMgeOO4/exec"
    print(f"[*] Pushing to Google Cloud...")
    import json, urllib.request, subprocess
    try:
        req = urllib.request.Request(gas_url, data=json.dumps(lesson_plan).encode("utf-8"), headers={"Content-Type": "application/json"})
        with urllib.request.urlopen(req, timeout=15) as resp:
            res = json.loads(resp.read().decode("utf-8"))
            if res.get("success") and res.get("url"):
                print(f"[!!!] GOOGLE DOC: {res.get("url")}")
                subprocess.run(["open", res.get("url")])
    except Exception as e: print(f"[!] Cloud Connection Failed: {e}")
    # ----------------------------------------

    for session in lesson_plan.get("sessions", []):
        document = Document(str(template_path))
        _apply_base_styles(document, config)
        _render_session_doc(document, session, config)
        session_output = output_dir / session["output_filename"]
        document.save(str(session_output))
        rendered_paths.append(session_output)

    return rendered_paths


def _render_session_markdown(session: dict[str, Any], config: dict[str, Any]) -> list[str]:
    view = _build_compact_session_view(session, config)
    lines = [
        f"# {view['title_line']}",
    ]
    if view["iep_students_line"]:
        lines.extend(["", view["iep_students_line"]])
    lines.extend(
        [
            "",
            "## DO NOW",
            view["do_now"],
            "",
            "## STANDARDS",
            "",
        ]
    )
    lines.extend(_markdown_table(view["standards_headers"], [view["standards_cells"]]))
    lines.extend(
        [
            "",
            "## OBJECTIVES",
            f"**Content Objective:** {view['content_objective']}",
            view["checkbox_line"],
            "",
            f"**Language Objective:** {view['language_objective']}",
            view["checkbox_line"],
            "",
            "## LESSON PROCEDURES",
            "",
        ]
    )
    lines.extend(
        _markdown_table(
            ["Phase (Time)", "Teacher Moves", "Student Moves", "SPED Supports", "ESOL/WIDA Scaffolds", "Formative Check"],
            [
                [
                    row["phase_time"],
                    _join_markdown_lines(row["teacher_moves"]),
                    _join_markdown_lines(row["student_moves"]),
                    _join_markdown_lines(row["sped_supports"]),
                    _join_markdown_lines(row["wida_scaffolds"]),
                    _join_markdown_lines(row["formative_check"]),
                ]
                for row in view["procedure_rows"]
            ],
        )
    )
    lines.extend(["", "## VOCABULARY", ""])
    lines.extend(
        _markdown_table(
            ["Tier", "Term", "Spanish", "Student-Friendly Definition", "Morphology", "EN-ES Cognate", "Cross-Disciplinary"],
            [
                [
                    row["tier"],
                    row["term"],
                    row["spanish"],
                    row["definition"],
                    row["morphology"],
                    row["cognate"],
                    row["cross_disciplinary"],
                ]
                for row in view["vocabulary_rows"]
            ],
        )
    )
    if view["accommodations_matrix_rows"]:
        lines.extend(["", "## IEP ACCOMMODATIONS MATRIX", ""])
        lines.extend(
            _markdown_table(
                ["Student", "Active Accommodations This Lesson"],
                [[row["student"], row["supports"]] for row in view["accommodations_matrix_rows"]],
            )
        )
    return lines


def _render_session_doc(document: Document, session: dict[str, Any], config: dict[str, Any]) -> None:
    view = _build_compact_session_view(session, config)

    title = document.add_paragraph(style="Normal")
    title_run = title.add_run(view["teacher_name"])
    title_run.bold = True
    title_run.font.size = Pt(14)
    title_tail = title.add_run(f" | {view['title_body']} | {view['display_date']}")
    title_tail.font.size = Pt(14)
    if view["iep_students_line"]:
        _add_paragraph(document, view["iep_students_line"])

    _add_heading(document, "DO NOW")
    _add_paragraph(document, view["do_now"])

    _add_heading(document, "STANDARDS")
    standards_table = document.add_table(rows=2, cols=3)
    standards_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _write_header_row(standards_table.rows[0].cells, view["standards_headers"])
    for index, text in enumerate(view["standards_cells"]):
        _write_cell_lines(standards_table.rows[1].cells[index], [text])

    _add_heading(document, "OBJECTIVES")
    _add_objective_line(document, "Content Objective:", view["content_objective"])
    _add_paragraph(document, view["checkbox_line"])
    _add_objective_line(document, "Language Objective:", view["language_objective"])
    _add_paragraph(document, view["checkbox_line"])

    _add_heading(document, "LESSON PROCEDURES")
    procedures_table = document.add_table(rows=1, cols=6)
    procedures_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _write_header_row(
        procedures_table.rows[0].cells,
        ["Phase (Time)", "Teacher Moves", "Student Moves", "SPED Supports", "ESOL/WIDA Scaffolds", "Formative Check"],
    )
    for row in view["procedure_rows"]:
        cells = procedures_table.add_row().cells
        _write_cell_lines(cells[0], [row["phase_time"]])
        _write_bullet_cell(cells[1], row["teacher_moves"])
        _write_bullet_cell(cells[2], row["student_moves"])
        _write_cell_lines(cells[3], row["sped_supports"])
        _write_cell_lines(cells[4], row["wida_scaffolds"])
        _write_cell_lines(cells[5], row["formative_check"])

    _add_heading(document, "VOCABULARY")
    vocab_table = document.add_table(rows=1, cols=7)
    vocab_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _write_header_row(
        vocab_table.rows[0].cells,
        ["Tier", "Term", "Spanish", "Student-Friendly Definition", "Morphology", "EN-ES Cognate", "Cross-Disciplinary"],
    )
    for row in view["vocabulary_rows"]:
        cells = vocab_table.add_row().cells
        _write_cell_lines(cells[0], [row["tier"]])
        _write_cell_lines(cells[1], [row["term"]])
        _write_cell_lines(cells[2], [row["spanish"]])
        _write_cell_lines(cells[3], [row["definition"]])
        _write_cell_lines(cells[4], [row["morphology"]])
        _write_cell_lines(cells[5], [row["cognate"]])
        _write_cell_lines(cells[6], [row["cross_disciplinary"]])

    if view["accommodations_matrix_rows"]:
        _add_heading(document, "IEP ACCOMMODATIONS MATRIX")
        matrix_table = document.add_table(rows=1, cols=2)
        matrix_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        _write_header_row(matrix_table.rows[0].cells, ["Student", "Active Accommodations This Lesson"])
        for row in view["accommodations_matrix_rows"]:
            cells = matrix_table.add_row().cells
            _write_cell_lines(cells[0], [row["student"]])
            _write_cell_lines(cells[1], [row["supports"]])


def _build_compact_session_view(session: dict[str, Any], config: dict[str, Any]) -> dict[str, Any]:
    info = session["lesson_information"]
    context = session.get("reference_render_context", {})
    supports = session["differentiation_sped_esol_supports_and_teacher_notes"]

    teacher_name = clean_line(str(config.get("teacher_name", "") or "Teacher Name"))
    display_date = _display_date(info["date"])
    title_body = _title_body(info["lesson_title"], info["session_label"])
    vocabulary_rows = _build_vocabulary_rows(session)
    procedure_rows = _build_procedure_rows(session, supports, vocabulary_rows)
    content_objective = _build_content_objective(session)
    iep_students_line = _build_iep_students_line(session)
    accommodations_matrix_rows = _build_accommodations_matrix_rows(session)

    return {
        "teacher_name": teacher_name,
        "title_body": title_body,
        "display_date": display_date,
        "title_line": f"{teacher_name} | {title_body} | {display_date}",
        "iep_students_line": iep_students_line,
        "do_now": _build_do_now_prompt(session),
        "standards_headers": [
            f"Grade {_extract_grade_number(info['course_or_grade'], config) - 1 if _extract_grade_number(info['course_or_grade'], config) > 1 else 1} Prerequisite",
            f"Grade {_extract_grade_number(info['course_or_grade'], config)} Target Standard",
            f"Grade {_extract_grade_number(info['course_or_grade'], config) + 1} Extension",
        ],
        "standards_cells": _build_standards_cells(session, config),
        "content_objective": content_objective,
        "language_objective": _build_language_objective(session, vocabulary_rows),
        "checkbox_line": "☐ Before | ☐ After",
        "procedure_rows": procedure_rows,
        "vocabulary_rows": vocabulary_rows,
        "accommodations_matrix_rows": accommodations_matrix_rows,
        "lesson_topic": clean_line(context.get("lesson_topic") or info["lesson_title"]),
    }


def _display_date(value: str) -> str:
    cleaned = clean_line(value)
    try:
        parsed = datetime.strptime(cleaned, "%Y-%m-%d")
        return parsed.strftime("%B %d, %Y").replace(" 0", " ")
    except ValueError:
        return cleaned


def _title_body(lesson_title: str, session_label: str) -> str:
    title = clean_line(lesson_title)
    session = clean_line(session_label)
    if session.lower() in title.lower():
        return title
    return f"{title} {session}"


def _build_do_now_prompt(session: dict[str, Any]) -> str:
    opening = session["opening_warm_up_launch"]
    context = session.get("reference_render_context", {})
    prompt_lines = _prompt_candidates(
        context.get("be_curious_prompts", []),
        opening.get("focus_tasks", []),
        _split_source_excerpt(opening.get("source_excerpt", "")),
    )
    combined = _combine_prompt_lines(prompt_lines, limit=2)
    if combined:
        return combined
    return _first_nonempty(
        opening.get("focus_tasks", []),
        opening.get("source_excerpt", ""),
        session["guided_practice_collaborative_learning"].get("source_excerpt", ""),
    )


def _build_iep_students_line(session: dict[str, Any]) -> str:
    supports = session.get("differentiation_sped_esol_supports_and_teacher_notes", {}).get("sped", [])
    labels = [clean_line(item.get("student", "")) for item in supports if clean_line(item.get("student", ""))]
    if not labels:
        return ""
    return "IEP Students: " + " | ".join(labels)


def _build_accommodations_matrix_rows(session: dict[str, Any]) -> list[dict[str, str]]:
    supports = session.get("differentiation_sped_esol_supports_and_teacher_notes", {}).get("sped", [])
    rows: list[dict[str, str]] = []
    for item in supports:
        student = clean_line(item.get("student", ""))
        matrix_supports = clean_line(item.get("matrix_supports", ""))
        if not student:
            continue
        if not matrix_supports:
            matrix_supports = ", ".join(clean_line(value) for value in item.get("supports", []) if clean_line(value))
        rows.append({"student": student, "supports": matrix_supports})
    return rows


def _build_standards_cells(session: dict[str, Any], config: dict[str, Any]) -> list[str]:
    grade = _extract_grade_number(session["lesson_information"]["course_or_grade"], config)
    standards = session["standards_and_learning_targets"]
    current_standards = standards.get("standards", [])
    current_standard = clean_line(current_standards[0]) if current_standards else ""
    current_cell = _format_target_standard_cell(session, standards)
    prerequisite_code = STANDARD_ALIGNMENT_MAP.get(current_standard, {}).get("prior", "")
    extension_code = STANDARD_ALIGNMENT_MAP.get(current_standard, {}).get("next", "")
    return [
        _format_standard_cell(prerequisite_code) or _build_prerequisite_descriptor(session, grade),
        current_cell,
        _format_standard_cell(extension_code) or _build_extension_descriptor(session, grade),
    ]


def _build_prerequisite_descriptor(session: dict[str, Any], grade: int) -> str:
    blob = _combined_source_blob(session)
    if "area" in blob:
        return f"Grade {max(grade - 1, 1)} prerequisite reasoning: Review rectangle area, unit-square reasoning, and how base and height describe a figure."
    if "independent variable" in blob or "dependent variable" in blob:
        return f"Grade {max(grade - 1, 1)} prerequisite reasoning: Review input/output pairs, patterns, and how one quantity changes with another."
    if "equation" in blob or "rate" in blob:
        return f"Grade {max(grade - 1, 1)} prerequisite reasoning: Review patterns, repeated addition, and how a rule generates outputs from inputs."
    return f"Grade {max(grade - 1, 1)} prerequisite reasoning: Review the prior vocabulary, model, and strategy students need before today's source task."


def _build_extension_descriptor(session: dict[str, Any], grade: int) -> str:
    blob = _combined_source_blob(session)
    if "area" in blob:
        return f"Grade {grade + 1} extension: Extend the strategy to justify formulas, compare composite figures, and solve for missing dimensions."
    if "independent variable" in blob or "dependent variable" in blob:
        return f"Grade {grade + 1} extension: Extend the lesson to graph relationships, compare rates, and justify how the equation matches the table."
    if "equation" in blob or "rate" in blob:
        return f"Grade {grade + 1} extension: Extend the lesson to proportional relationships, graph interpretation, and multi-step reasoning."
    return f"Grade {grade + 1} extension: Extend the source strategy into stronger explanation, transfer, and independent problem solving."


def _build_content_objective(session: dict[str, Any]) -> str:
    standards = session["standards_and_learning_targets"]
    lesson_objective = clean_line(session["lesson_objective_and_student_success_criteria"]["lesson_objective"]).rstrip(".")
    candidate = _preferred_learning_target(session, standards)
    if candidate:
        return _ensure_objective_voice(candidate, "I can")
    if lesson_objective.lower().startswith("i can "):
        return lesson_objective
    if lesson_objective:
        return _ensure_objective_voice(f"I can {lesson_objective}", "I can")
    return "I can explain the lesson goal using the source task and vocabulary."


def _build_language_objective(session: dict[str, Any], vocabulary_rows: list[dict[str, str]]) -> str:
    vocab_terms = [row["term"].lower() for row in vocabulary_rows[:3]]
    preferred_target = _preferred_learning_target(session, session["standards_and_learning_targets"]).lower()
    blob = f"{_combined_source_blob(session)} {preferred_target}"
    if "independent variable" in blob or "dependent variable" in blob:
        focus = "proportional relationships"
    elif any(term in blob for term in ("composite figure", "irregular shape", "trapezoid", "flag", "swallowtail")):
        focus = "how a composite or irregular figure can be decomposed into shapes with known area"
    elif "area" in blob and any(term in blob for term in ("polygon", "octagon", "rhombus")):
        focus = "how a polygon can be decomposed into familiar shapes to find its area"
    elif "area" in blob and any(term in blob for term in ("triangle", "base", "height")):
        focus = "how the area of a triangle is related to a rectangle or parallelogram"
    else:
        focus = "my reasoning about the lesson task"
    vocab_phrase = _natural_list(vocab_terms) if vocab_terms else "the lesson vocabulary"
    return f"I will explain {focus} using sentence frames with the terms {vocab_phrase}."


def _format_target_standard_cell(session: dict[str, Any], standards: dict[str, Any]) -> str:
    current_standards = [clean_line(item) for item in standards.get("standards", []) if clean_line(item)]
    if current_standards:
        current_standard = current_standards[0]
        learning_target = _preferred_learning_target(session, standards).rstrip(".")
        description = (
            _source_standard_description(standards, current_standard)
            or STANDARD_DESCRIPTION_OVERRIDES.get(current_standard, "")
            or learning_target
        )
        return _format_standard_cell(current_standard, description)
    preferred_target = _preferred_learning_target(session, standards)
    if preferred_target:
        return f"Source learning target: {preferred_target}"
    return standards.get("standards_status", "Not explicitly listed in source slides.")


def _preferred_learning_target(session: dict[str, Any], standards: dict[str, Any]) -> str:
    lesson_objective = clean_line(
        session.get("lesson_objective_and_student_success_criteria", {}).get("lesson_objective", "")
    ).rstrip(".")
    for candidate in standards.get("i_can_statements", []) or standards.get("learning_targets", []):
        cleaned_candidate = clean_line(candidate).rstrip(".")
        stripped_candidate = re.sub(r"^(?:I can|We will)\s+", "", cleaned_candidate, flags=re.IGNORECASE).rstrip(".")
        if lesson_objective and lesson_objective.lower() in {cleaned_candidate.lower(), stripped_candidate.lower()}:
            return cleaned_candidate
    return select_preferred_learning_target(
        standards.get("i_can_statements", []) or standards.get("learning_targets", []),
        _combined_source_blob(session),
    ) or _first_nonempty(standards.get("i_can_statements", []), standards.get("learning_targets", []))


def _format_standard_cell(code: str, description: str = "") -> str:
    normalized_code = clean_line(code)
    if not normalized_code:
        return ""
    normalized_description = clean_line(description or STANDARD_DESCRIPTION_OVERRIDES.get(normalized_code, ""))
    if normalized_description:
        return f"{normalized_code}: {normalized_description}"
    return normalized_code


def _source_standard_description(standards: dict[str, Any], code: str) -> str:
    source_lines = [clean_line(item) for item in standards.get("standards_source", {}).get("source_lines", [])]
    for line in source_lines:
        if code.lower() not in line.lower():
            continue
        cleaned = re.sub(r"^standard[s]?:?\s*", "", line, flags=re.IGNORECASE)
        cleaned = re.sub(rf"^{re.escape(code)}\s*[:\-]?\s*", "", cleaned, flags=re.IGNORECASE)
        cleaned = clean_line(cleaned)
        if cleaned and cleaned.lower() != code.lower():
            return cleaned
    return ""


def _ensure_objective_voice(text: str, lead: str) -> str:
    cleaned = clean_line(text).rstrip(".")
    if not cleaned:
        return cleaned
    lowered = cleaned.lower()
    lead_lower = lead.lower()
    if lowered.startswith(f"{lead_lower} "):
        return cleaned
    return f"{lead} {cleaned[:1].lower() + cleaned[1:]}"


def _natural_list(items: list[str]) -> str:
    cleaned = [clean_line(item) for item in items if clean_line(item)]
    if not cleaned:
        return ""
    if len(cleaned) == 1:
        return cleaned[0]
    if len(cleaned) == 2:
        return f"{cleaned[0]} and {cleaned[1]}"
    return ", ".join(cleaned[:-1]) + f", and {cleaned[-1]}"


def _prompt_candidates(*groups: Iterable[str] | str) -> list[str]:
    results: list[str] = []
    seen: set[str] = set()
    for group in groups:
        if isinstance(group, str):
            items = [group]
        else:
            items = list(group)
        for item in items:
            normalized = _normalize_prompt_fragment(item)
            if not normalized:
                continue
            lowered = normalized.lower()
            if lowered in seen:
                continue
            seen.add(lowered)
            results.append(normalized)
    return results


def _normalize_prompt_fragment(value: str) -> str:
    candidate = clean_line(str(value))
    if not candidate:
        return ""
    candidate = _strip_phase_prefix(candidate)
    if not candidate:
        return ""
    lowered = candidate.lower()
    if lowered in GENERIC_PHASE_LINES:
        return ""
    return candidate


def _split_source_excerpt(value: str) -> list[str]:
    cleaned = clean_line(value)
    if not cleaned:
        return []
    cleaned = _strip_phase_prefix(cleaned)
    parts = [clean_line(part) for part in re.split(r"(?<=[?!])\s+|(?<=\.)\s+", cleaned) if clean_line(part)]
    return parts or [cleaned]


def _strip_phase_prefix(value: str) -> str:
    cleaned = clean_line(value)
    if not cleaned:
        return ""
    prefixes = [
        "Be Curious",
        "Closure",
        "Discussion",
        "Do Now",
        "Exit Ticket",
        "Guided Practice",
        "Independent Practice",
        "Launch",
        "Learning Target",
        "Lesson launch",
        "Mini-Lesson",
        "Model",
        "Session 1",
        "Session 2",
        "Warm-Up",
        "Warm Up",
        "Workspace",
    ]
    result = cleaned
    changed = True
    while changed:
        changed = False
        for prefix in prefixes:
            pattern = rf"^{re.escape(prefix)}(?:\s*[:\-])?\s+"
            updated = re.sub(pattern, "", result, flags=re.IGNORECASE)
            if updated != result:
                result = clean_line(updated)
                changed = True
    return result


def _combine_prompt_lines(lines: list[str], limit: int = 2) -> str:
    selected = [clean_line(line) for line in lines[:limit] if clean_line(line)]
    if not selected:
        return ""
    combined = " ".join(selected)
    return clean_line(combined)



def _build_procedure_rows(
    session: dict[str, Any],
    supports: dict[str, Any],
    vocabulary_rows: list[dict[str, str]],
) -> list[dict[str, Any]]:
    opening = session["opening_warm_up_launch"]
    modeling = session["mini_lesson_modeling_concept_development"]
    guided = session["guided_practice_collaborative_learning"]
    independent = session["independent_practice_application_stations"]
    closure = session["closure_exit_ticket_assessment"]
    context = session.get("reference_render_context", {})

    opening_times = _split_time(opening["time_minutes"], [0.35, 0.3, 0.35])
    closure_times = _split_time(closure["time_minutes"], [0.45, 0.55])

    row_payloads = {
        "warm_up": {
            "time": opening_times[0],
            "teacher": _prepend_prompt(opening["teacher_actions"][:2], opening["source_excerpt"]),
            "student": opening["student_actions"][:2] or opening["focus_tasks"][:2],
            "source": opening["source_excerpt"],
            "evidence": opening["evidence_of_learning"][:2],
        },
        "launch": {
            "time": opening_times[1],
            "teacher": _build_launch_teacher_moves(session),
            "student": _build_launch_student_moves(session),
            "source": _combine_prompt_lines(
                _prompt_candidates(
                    opening.get("focus_tasks", []),
                    _split_source_excerpt(opening.get("source_excerpt", "")),
                ),
                limit=2,
            ),
            "evidence": opening["evidence_of_learning"][:1] + objective_success_lookfors(session),
        },
        "be_curious": {
            "time": opening_times[2],
            "teacher": _build_be_curious_teacher_moves(session),
            "student": _build_be_curious_student_moves(session),
            "source": _combine_prompt_lines(
                _prompt_candidates(
                    context.get("be_curious_prompts", []),
                    _split_source_excerpt(opening.get("source_excerpt", "")),
                ),
                limit=2,
            ),
            "evidence": ["Collect one source-based noticing or wondering before moving into formal practice."],
        },
        "guided_practice": {
            "time": modeling["time_minutes"],
            "teacher": _build_guided_teacher_moves(session),
            "student": _build_guided_student_moves(session),
            "source": modeling["source_excerpt"],
            "evidence": modeling["evidence_of_learning"][:2],
        },
        "collaborative_practice": {
            "time": guided["time_minutes"],
            "teacher": _build_collaborative_teacher_moves(session),
            "student": _build_collaborative_student_moves(session),
            "source": guided["source_excerpt"],
            "evidence": guided["evidence_of_learning"][:2],
        },
        "independent_practice": {
            "time": independent["time_minutes"],
            "teacher": _build_independent_teacher_moves(session),
            "student": _build_independent_student_moves(session),
            "source": independent["source_excerpt"],
            "evidence": independent["evidence_of_learning"][:2],
        },
        "closing": {
            "time": closure_times[0],
            "teacher": _build_closing_teacher_moves(closure),
            "student": _build_closing_student_moves(session),
            "source": closure["source_excerpt"],
            "evidence": closure["evidence_of_learning"][:1],
        },
        "exit_ticket": {
            "time": closure_times[1],
            "teacher": _build_exit_teacher_moves(closure),
            "student": _build_exit_student_moves(closure, independent),
            "source": _first_nonempty(closure.get("focus_tasks", []), closure["source_excerpt"]),
            "evidence": closure["evidence_of_learning"][:2] or independent["evidence_of_learning"][:1],
        },
    }

    rows = []
    for key, label in PROCEDURE_ROW_ORDER:
        payload = row_payloads[key]
        rows.append(
            {
                "phase_time": f"{label}\n({payload['time']} min)",
                "teacher_moves": _normalize_procedure_lines(payload["teacher"], limit=_teacher_move_limit(key)),
                "student_moves": _normalize_procedure_lines(payload["student"], limit=_student_move_limit(key)),
                "sped_supports": _build_sped_support_lines(key, supports),
                "wida_scaffolds": _build_wida_scaffolds(key, payload["source"], vocabulary_rows),
                "formative_check": _build_formative_check_lines(key, payload["evidence"], payload["source"]),
            }
        )
    return rows


def objective_success_lookfors(session: dict[str, Any]) -> list[str]:
    return [
        f"Listen for a student explanation connected to this success criterion: {item}"
        for item in session["lesson_objective_and_student_success_criteria"].get("student_success_criteria", [])[:1]
    ]


def _build_launch_teacher_moves(session: dict[str, Any]) -> list[str]:
    opening_prompt = _build_do_now_prompt(session)
    content_objective = _build_content_objective(session)
    guided_prompt = _first_nonempty(
        session["guided_practice_collaborative_learning"].get("focus_tasks", []),
        session["guided_practice_collaborative_learning"].get("source_excerpt", ""),
    )
    return [
        f"Preview the learning target: {content_objective}",
        f"Show the opening task: {opening_prompt}",
        f"Bridge the opening task to the first worked example or guided problem: {truncate(guided_prompt, 88)}",
    ]


def _build_launch_student_moves(session: dict[str, Any]) -> list[str]:
    return [
        "Read the learning target and connect it to the opening task.",
        "Respond to the opening task and record a first strategy, estimate, or observation.",
        "Share one prediction about how the lesson problem will be solved.",
    ]


def _build_be_curious_teacher_moves(session: dict[str, Any]) -> list[str]:
    context = session.get("reference_render_context", {})
    opening = session["opening_warm_up_launch"]
    return [
        f"Display the notice/wonder prompt: {_combine_prompt_lines(_prompt_candidates(context.get('be_curious_prompts', []), _split_source_excerpt(opening['source_excerpt'])), limit=2)}",
        "Give students quick partner talk time for one noticing and one wondering.",
        "Record a few student ideas before formal modeling begins.",
    ]


def _build_be_curious_student_moves(session: dict[str, Any]) -> list[str]:
    context = session.get("reference_render_context", {})
    prompt = _combine_prompt_lines(
        _prompt_candidates(
            context.get("be_curious_prompts", []),
            _split_source_excerpt(session["opening_warm_up_launch"]["source_excerpt"]),
        ),
        limit=2,
    )
    return [
        f"Respond to the prompt: {prompt}",
        "Share one noticing and one wondering with a partner.",
        "Record one sentence stem that connects the observation to the lesson math.",
    ]


def _build_guided_teacher_moves(session: dict[str, Any]) -> list[str]:
    modeling = session["mini_lesson_modeling_concept_development"]
    context = session.get("reference_render_context", {})
    guided = session["guided_practice_collaborative_learning"]
    moves = list(modeling.get("teacher_actions", []))
    worked_examples = context.get("worked_examples", [])
    if worked_examples:
        prompt = clean_line(str(worked_examples[0].get("prompt", "")))
        reveal = clean_line(str(worked_examples[0].get("reveal", "")))
        if prompt:
            moves.append(f"Work through the model prompt: {prompt}")
        if reveal and reveal.lower() != prompt.lower():
            moves.append(f"Think aloud with the revealed relationship or example: {reveal}")
    if context.get("guided_practice"):
        moves.append(f"Release students to solve the next example: {truncate(context['guided_practice'][0], 140)}")
    if guided.get("teacher_actions") and clean_line(guided["teacher_actions"][-1]):
        moves.append(guided["teacher_actions"][-1])
    return moves


def _build_guided_student_moves(session: dict[str, Any]) -> list[str]:
    modeling = session["mini_lesson_modeling_concept_development"]
    context = session.get("reference_render_context", {})
    moves = list(modeling.get("student_actions", []))
    if context.get("guided_practice"):
        moves.append(f"Solve the guided example from the slides: {truncate(context['guided_practice'][0], 140)}")
    return moves


def _build_collaborative_teacher_moves(session: dict[str, Any]) -> list[str]:
    guided = session["guided_practice_collaborative_learning"]
    context = session.get("reference_render_context", {})
    moves = list(guided.get("teacher_actions", []))
    collaborative_prompt = _optional_first_nonempty(context.get("collaborative_tasks", []))
    if collaborative_prompt:
        moves.append(f"Keep partner talk anchored to the discussion task: {truncate(collaborative_prompt, 140)}")
    closing_check = _optional_first_nonempty(context.get("checks_for_understanding", []))
    if closing_check:
        moves.append(f"Use the discussion check to close the task: {truncate(closing_check, 140)}")
    small_group_students = _small_group_students(session)
    if small_group_students:
        small_group_prompt = _small_group_task_prompt(session, "collaborative_practice")
        moves.append(
            "Pull "
            + _natural_list(small_group_students)
            + " for a brief small group on the same slide/book discussion: "
            + truncate(small_group_prompt, 140)
        )
    return moves


def _build_collaborative_student_moves(session: dict[str, Any]) -> list[str]:
    guided = session["guided_practice_collaborative_learning"]
    context = session.get("reference_render_context", {})
    moves = list(guided.get("student_actions", []))
    collaborative_prompt = _optional_first_nonempty(context.get("collaborative_tasks", []))
    if collaborative_prompt:
        moves.append(f"Discuss the collaborative prompt: {truncate(collaborative_prompt, 140)}")
    closing_check = _optional_first_nonempty(context.get("checks_for_understanding", []))
    if closing_check:
        moves.append(f"Record a response to the check-for-understanding prompt: {truncate(closing_check, 140)}")
    if _small_group_students(session):
        small_group_prompt = _small_group_task_prompt(session, "collaborative_practice")
        moves.append(
            "In the small group, rehearse the same slide/book discussion with frames and the visual: "
            + truncate(small_group_prompt, 140)
        )
    return moves


def _build_independent_teacher_moves(session: dict[str, Any]) -> list[str]:
    independent = session["independent_practice_application_stations"]
    context = session.get("reference_render_context", {})
    moves = list(independent.get("teacher_actions", []))
    independent_tasks = context.get("independent_practice", [])
    follow_through = clean_line(independent_tasks[1]) if len(independent_tasks) > 1 else ""
    if follow_through:
        moves.append(f"Require students to complete the follow-through direction: {truncate(follow_through, 140)}")
    follow_up_task = _optional_first_nonempty(context.get("lets_explore_more_tasks", []))
    if follow_up_task:
        moves.append(f"Use the follow-up task only if it appears in the slides: {truncate(follow_up_task, 140)}")
    if context.get("reveal_math_workbook_references"):
        moves.append("Use the Reveal Math Workspace or workbook task exactly as referenced in the lesson.")
    small_group_students = _small_group_students(session)
    if small_group_students:
        small_group_prompt = _small_group_task_prompt(session, "independent_practice")
        moves.append(
            "Meet with "
            + _natural_list(small_group_students)
            + " in a brief small group to chunk the same slide/book task before releasing them back to independent work: "
            + truncate(small_group_prompt, 140)
        )
    return moves


def _build_independent_student_moves(session: dict[str, Any]) -> list[str]:
    independent = session["independent_practice_application_stations"]
    context = session.get("reference_render_context", {})
    moves = list(independent.get("student_actions", []))
    independent_tasks = context.get("independent_practice", [])
    prompt = _optional_first_nonempty(independent_tasks)
    if prompt:
        moves.append(f"Complete the independent prompt: {truncate(prompt, 140)}")
    follow_through = clean_line(independent_tasks[1]) if len(independent_tasks) > 1 else ""
    if follow_through:
        moves.append(f"Follow the source direction exactly: {truncate(follow_through, 140)}")
    if _small_group_students(session):
        small_group_prompt = _small_group_task_prompt(session, "independent_practice")
        moves.append(
            "In the small group, solve the same slide/book task with chunked directions and highlighted quantities: "
            + truncate(small_group_prompt, 140)
        )
    return moves


def _build_closing_teacher_moves(closure: dict[str, Any]) -> list[str]:
    closing_prompt = _combine_prompt_lines(_prompt_candidates(closure.get("focus_tasks", []), closure.get("source_excerpt", "")), limit=1)
    return [
        f"Return to the closing prompt: {closing_prompt}",
        "Ask students to name the strategy or idea they need to keep practicing.",
        "Listen for precise mathematical language before dismissing the class or assigning follow-up work.",
    ]


def _build_closing_student_moves(session: dict[str, Any]) -> list[str]:
    closure = session["closure_exit_ticket_assessment"]
    context = session.get("reference_render_context", {})
    moves = list(closure.get("student_actions", []))
    if context.get("summary_closure_language"):
        moves.append(f"Respond to the closing language from the slides: {truncate(context['summary_closure_language'][0], 140)}")
    elif context.get("checks_for_understanding"):
        moves.append(f"Self-assess with the final prompt: {truncate(context['checks_for_understanding'][0], 140)}")
    return moves


def _build_exit_teacher_moves(closure: dict[str, Any]) -> list[str]:
    prompt = _combine_prompt_lines(_prompt_candidates(closure.get("focus_tasks", []), closure.get("source_excerpt", "")), limit=1)
    return [
        f"Assign the final prompt and require visible reasoning: {prompt}",
        "Set the expectation that students answer independently and show all work.",
        "Collect the written responses and sort them for the next grouping decision.",
    ]


def _build_exit_student_moves(closure: dict[str, Any], independent: dict[str, Any]) -> list[str]:
    prompt = _combine_prompt_lines(
        _prompt_candidates(
            closure.get("focus_tasks", []),
            independent.get("focus_tasks", []),
            closure.get("source_excerpt", ""),
        ),
        limit=1,
    )
    return [
        f"Complete the final check independently: {prompt}",
        "Show the strategy and label the answer clearly.",
        "Use the lesson vocabulary in at least one sentence of explanation.",
    ]


def _small_group_students(session: dict[str, Any]) -> list[str]:
    supports = session.get("differentiation_sped_esol_supports_and_teacher_notes", {}).get("sped", [])
    return [
        clean_line(item.get("student", ""))
        for item in supports
        if clean_line(item.get("student", "")) and student_requires_small_group(item)
    ]


def _small_group_task_prompt(session: dict[str, Any], row_kind: str) -> str:
    context = session.get("reference_render_context", {})
    if row_kind == "collaborative_practice":
        return _optional_first_nonempty(
            context.get("collaborative_tasks", []),
            context.get("guided_practice", []),
            session["guided_practice_collaborative_learning"].get("focus_tasks", []),
            session["guided_practice_collaborative_learning"].get("source_excerpt", ""),
        )
    return _optional_first_nonempty(
        context.get("independent_practice", []),
        context.get("lets_explore_more_tasks", []),
        context.get("reveal_math_workbook_references", []),
        session["independent_practice_application_stations"].get("focus_tasks", []),
        session["independent_practice_application_stations"].get("source_excerpt", ""),
    )


def _teacher_move_limit(row_kind: str) -> int:
    return {
        "warm_up": 3,
        "launch": 4,
        "be_curious": 4,
        "guided_practice": 5,
        "collaborative_practice": 6,
        "independent_practice": 6,
        "closing": 4,
        "exit_ticket": 4,
    }[row_kind]


def _student_move_limit(row_kind: str) -> int:
    return {
        "warm_up": 3,
        "launch": 4,
        "be_curious": 4,
        "guided_practice": 4,
        "collaborative_practice": 5,
        "independent_practice": 5,
        "closing": 3,
        "exit_ticket": 3,
    }[row_kind]


def _build_sped_support_lines(row_kind: str, supports: dict[str, Any]) -> list[str]:
    rows = []
    for item in supports.get("sped", []):
        selected = _select_supports_for_row(item, row_kind)
        if not selected:
            selected = item.get("supports", [])[:2]
        rows.append(f"{item['student']}: {', '.join(selected)}")
    return rows or ["Use the configured support profile for the students who need extra scaffolds."]


def _select_supports_for_row(item: dict[str, Any], row_kind: str) -> list[str]:
    normalized = [clean_line(value) for value in item.get("supports", []) if clean_line(value)]
    if row_kind in {"collaborative_practice", "independent_practice"} and student_requires_small_group(item):
        normalized.append("small group")
    normalized = unique_preserve(normalized)
    priorities = SUPPORT_PRIORITIES[row_kind]
    selected = [item for item in normalized if item in priorities]
    ordered = [item for item in priorities if item in selected]
    return ordered[:3] if ordered else normalized[:2]


def _build_wida_scaffolds(row_kind: str, source_text: str, vocabulary_rows: list[dict[str, str]]) -> list[str]:
    vocab = vocabulary_rows[0]["term"].lower() if vocabulary_rows else "the lesson vocabulary"
    prompt = _combine_prompt_lines(_prompt_candidates(source_text), limit=2) or clean_line(source_text)
    if row_kind in {"warm_up", "launch", "be_curious"}:
        return [
            f"L1-2: Point to the key visual or quantity in this prompt: {truncate(prompt, 70)}",
            f"L3-4: Complete a sentence frame with {vocab}.",
            "L5: Explain the relationship in a full sentence using because.",
        ]
    if row_kind in {"guided_practice", "collaborative_practice"}:
        return [
            "L1-2: Use gestures, labels, or a short phrase while following the model.",
            f"L3-4: State the next step aloud with {vocab}.",
            "L5: Justify why the strategy works using a complete explanation.",
        ]
    if row_kind == "independent_practice":
        return [
            "L1-2: Highlight the numbers and identify the input/output or the base/height before solving.",
            "L3-4: Use a sentence frame to explain the answer.",
            "L5: Write a full explanation with transitions and precise vocabulary.",
        ]
    return [
        "L1-2: Point to one checklist item or answer choice.",
        f"L3-4: State one reflection using {vocab}.",
        "L5: Explain why the answer is correct and what still needs practice.",
    ]


def _build_formative_check_lines(row_kind: str, evidence_lines: list[str], source_text: str) -> list[str]:
    if evidence_lines:
        details = [_clean_procedure_line(item) or clean_line(item) for item in evidence_lines if clean_line(item)]
        if details:
            return details[:2]
    prompt = _combine_prompt_lines(_prompt_candidates(source_text), limit=2) or clean_line(source_text)
    if row_kind in {"warm_up", "launch", "be_curious"}:
        return [f"Listen for a quick lesson-based response to: {truncate(prompt, 80)}"]
    if row_kind in {"guided_practice", "collaborative_practice"}:
        return ["Check whether students can explain the strategy, not just give the answer."]
    if row_kind == "independent_practice":
        return ["Scan written work for accurate setup, labels, and transfer of the modeled strategy."]
    return ["Collect the final response and use it to plan the next groupings or re-teach move."]


def _build_vocabulary_rows(session: dict[str, Any]) -> list[dict[str, str]]:
    terms = _extracted_vocabulary(session) or _fallback_vocabulary_items(session)
    rows = []
    seen: set[str] = set()
    for term, definition in terms:
        normalized = term.lower()
        if normalized in seen:
            continue
        seen.add(normalized)
        spanish = SPANISH_TRANSLATIONS.get(normalized, normalized)
        rows.append(
            {
                "tier": _tier_for_term(normalized),
                "term": term,
                "spanish": spanish,
                "definition": definition,
                "morphology": MORPHOLOGY_HINTS.get(normalized, _fallback_morphology(term)),
                "cognate": _cognate_label(term, spanish),
                "cross_disciplinary": CROSS_DISCIPLINARY_HINTS.get(normalized, _fallback_cross_disciplinary(session)),
            }
        )
        if len(rows) >= 5:
            break
    return rows or [
        {
            "tier": "2",
            "term": "Lesson Focus",
            "spanish": "enfoque de la leccion",
            "definition": f"The key mathematical idea students work on in {clean_line(session['lesson_information']['lesson_title']).lower()}.",
            "morphology": "lesson + focus",
            "cognate": "—",
            "cross_disciplinary": _fallback_cross_disciplinary(session),
        }
    ]


def _extracted_vocabulary(session: dict[str, Any]) -> list[tuple[str, str]]:
    context = session.get("reference_render_context", {})
    rows = []
    for item in context.get("vocabulary_terms", []):
        term = clean_line(str(item.get("term", "")))
        definition = clean_line(str(item.get("definition", "")))
        if term and definition:
            rows.append((term.title(), definition))
    return rows


def _fallback_vocabulary_items(session: dict[str, Any]) -> list[tuple[str, str]]:
    standards = session.get("standards_and_learning_targets", {})
    preferred_target = _preferred_learning_target(session, standards)
    search_blob = " ".join(
        piece
        for piece in (
            _combined_source_blob(session),
            clean_line(session["lesson_information"]["lesson_title"]),
            clean_line(preferred_target),
        )
        if clean_line(piece)
    ).lower()
    items = []
    for term, definition in DEFAULT_VOCAB_DEFINITIONS.items():
        if re.search(rf"\b{re.escape(term)}\b", search_blob):
            items.append((term.title(), definition))
    if len(items) < 3:
        for term in (
            "area",
            "composite figure",
            "trapezoid",
            "rhombus",
            "octagon",
            "polygon",
            "triangle",
            "decompose",
            "compose",
            "diagonal",
        ):
            if not re.search(rf"\b{re.escape(term)}\b", search_blob):
                continue
            entry = (term.title(), DEFAULT_VOCAB_DEFINITIONS[term])
            if entry not in items:
                items.append(entry)
            if len(items) >= 5:
                break
    return items


def _tier_for_term(term: str) -> str:
    if term in {"rate", "area", "base", "height", "equation", "variable", "compose", "decompose"}:
        return "2"
    if term in {"composite figure", "trapezoid", "rhombus", "octagon", "diagonal"}:
        return "3"
    if "relationship" in term or "proportionality" in term or "independent" in term or "dependent" in term:
        return "3"
    return "2"


def _fallback_morphology(term: str) -> str:
    parts = [piece for piece in re.split(r"[\s\-]+", term.lower()) if piece]
    return " + ".join(parts[:3]) if parts else "root"


def _cognate_label(term: str, spanish: str) -> str:
    normalized_term = term.lower().replace(" ", "")
    normalized_spanish = spanish.lower().replace(" ", "")
    if normalized_term[:4] == normalized_spanish[:4]:
        return f"✓ {spanish}"
    if normalized_term[:3] == normalized_spanish[:3]:
        return f"✓ {spanish}"
    return "—"


def _fallback_cross_disciplinary(session: dict[str, Any]) -> str:
    blob = _combined_source_blob(session)
    if "table" in blob or "graph" in blob:
        return "Science, Data"
    if "area" in blob or "polygon" in blob:
        return "Geometry, Design"
    return "Science, Problem Solving"


def _combined_source_blob(session: dict[str, Any]) -> str:
    pieces: list[str] = []
    for key in (
        "opening_warm_up_launch",
        "mini_lesson_modeling_concept_development",
        "guided_practice_collaborative_learning",
        "independent_practice_application_stations",
        "closure_exit_ticket_assessment",
    ):
        phase = session.get(key, {})
        pieces.extend(phase.get("focus_tasks", []))
        pieces.extend(phase.get("teacher_actions", []))
        pieces.extend(phase.get("student_actions", []))
        pieces.append(phase.get("source_excerpt", ""))
    context = session.get("reference_render_context", {})
    pieces.extend(context.get("reasoning_tasks", []))
    pieces.extend(item.get("term", "") for item in context.get("vocabulary_terms", []))
    return " ".join(clean_line(item).lower() for item in pieces if clean_line(item))


def _split_time(total: int, ratios: list[float]) -> list[int]:
    values = [max(int(round(total * ratio)), 1) for ratio in ratios]
    diff = total - sum(values)
    index = 0
    while diff != 0:
        values[index % len(values)] += 1 if diff > 0 else -1
        diff = total - sum(values)
        index += 1
    return values


def _prepend_prompt(items: list[str], source_excerpt: str) -> list[str]:
    cleaned = _normalize_procedure_lines(items, limit=4)
    if cleaned:
        return cleaned
    fallback = _normalize_prompt_fragment(source_excerpt)
    return [fallback] if fallback else []


def _normalize_procedure_lines(lines: list[str], limit: int = 4) -> list[str]:
    normalized: list[str] = []
    seen: set[str] = set()
    for line in lines:
        candidate = _clean_procedure_line(line)
        if not candidate:
            continue
        lowered = candidate.lower()
        if lowered in seen:
            continue
        seen.add(lowered)
        normalized.append(candidate)
        if len(normalized) >= limit:
            break
    return normalized


def _clean_procedure_line(value: str) -> str:
    candidate = clean_line(str(value))
    if not candidate:
        return ""
    substitutions = [
        (r"^Launch the actual opening from slides? [^:]+:\s*", "Display the opening prompt: "),
        (r"^Model the source sequence from slides? [^:]+:\s*", "Model: "),
        (r"^Coach the collaborative task from the source deck:\s*", ""),
        (r"^Assign the actual independent/application task:\s*", "Assign: "),
        (r"^Use the session close from the source slides:\s*", "Use the closing prompt: "),
        (r"^Use the source check question:\s*", "Ask: "),
        (r"^Press for the exact mathematical relationship named in the deck:\s*", "Press students to explain: "),
    ]
    for pattern, replacement in substitutions:
        candidate = re.sub(pattern, replacement, candidate, flags=re.IGNORECASE)
    candidate = re.sub(r"\bthe source deck\b", "the lesson", candidate, flags=re.IGNORECASE)
    candidate = re.sub(r"\bthe source slides\b", "the lesson", candidate, flags=re.IGNORECASE)
    candidate = re.sub(r"\bsource slide\b", "visual", candidate, flags=re.IGNORECASE)
    candidate = re.sub(r"\bsource-based\b", "lesson-based", candidate, flags=re.IGNORECASE)
    candidate = re.sub(r"\bactual\b", "", candidate, flags=re.IGNORECASE)
    candidate = re.sub(r"\bexact\b", "", candidate, flags=re.IGNORECASE)
    candidate = clean_line(candidate)
    if candidate.endswith(":"):
        return ""
    return candidate


def _first_nonempty(*groups: Iterable[str] | str) -> str:
    for group in groups:
        if isinstance(group, str):
            candidate = clean_line(group)
            if candidate:
                return candidate
            continue
        for item in group:
            candidate = clean_line(str(item))
            if candidate:
                return candidate
    return "Use the source task and explain the strategy clearly."


def _optional_first_nonempty(*groups: Iterable[str] | str) -> str:
    for group in groups:
        if isinstance(group, str):
            candidate = clean_line(group)
            if candidate:
                return candidate
            continue
        for item in group:
            candidate = clean_line(str(item))
            if candidate:
                return candidate
    return ""


def _extract_grade_number(course_or_grade: str, config: dict[str, Any]) -> int:
    match = re.search(r"\b(\d+)\b", clean_line(course_or_grade))
    if match:
        return max(int(match.group(1)), 1)
    return max(int(config.get("default_grade", 6)), 1)


def truncate(text: str, limit: int) -> str:
    cleaned = clean_line(text)
    if len(cleaned) <= limit:
        return cleaned
    return cleaned[: limit - 1].rstrip() + "…"


def _markdown_table(headers: list[str], rows: list[list[str]]) -> list[str]:
    formatted_rows = []
    for row in rows:
        formatted_rows.append(
            [
                clean_line(cell).replace("|", "\\|").replace("\n", "<br>")
                for cell in row
            ]
        )
    lines = [
        "| " + " | ".join(headers) + " |",
        "| " + " | ".join("---" for _ in headers) + " |",
    ]
    for row in formatted_rows:
        lines.append("| " + " | ".join(row) + " |")
    return lines


def _join_markdown_lines(lines: list[str]) -> str:
    return "<br>".join(clean_line(line) for line in lines if clean_line(line))


def _apply_base_styles(document: Document, config: dict[str, Any]) -> None:
    body_font = str((((config.get("docx_style") or {}).get("body_font")) or "Calibri"))
    heading_font = str((((config.get("docx_style") or {}).get("heading_font")) or body_font))

    normal_style = document.styles["Normal"]
    normal_style.font.name = body_font
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), body_font)
    normal_style.font.size = Pt(11)
    normal_style.paragraph_format.line_spacing = 1.15
    normal_style.paragraph_format.space_after = Pt(4)

    heading_style = document.styles["Heading 1"]
    heading_style.font.name = heading_font
    heading_style._element.rPr.rFonts.set(qn("w:eastAsia"), heading_font)
    heading_style.font.size = Pt(12)
    heading_style.font.bold = True


def _add_heading(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="Heading 1")
    paragraph.add_run(clean_line(text))


def _add_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="Normal")
    paragraph.add_run(clean_line(text))


def _add_objective_line(document: Document, label: str, text: str) -> None:
    paragraph = document.add_paragraph(style="Normal")
    run = paragraph.add_run(label + " ")
    run.bold = True
    paragraph.add_run(clean_line(text))


def _write_header_row(cells: list[Any], labels: list[str]) -> None:
    for index, label in enumerate(labels):
        _write_cell_lines(cells[index], [label], bold=True)


def _write_cell_lines(cell: Any, lines: list[str], *, bold: bool = False) -> None:
    cell.text = ""
    usable = [clean_line(line) for line in lines if clean_line(line)]
    if not usable:
        usable = [""]
    for index, line in enumerate(usable):
        paragraph = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(line)
        run.bold = bold


def _write_bullet_cell(cell: Any, lines: list[str]) -> None:
    items = [clean_line(line) for line in lines if clean_line(line)]
    if not items:
        _write_cell_lines(cell, [""])
        return
    cell.text = ""
    for index, item in enumerate(items):
        paragraph = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(f"• {item}")
