from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_PATH = (
    "/Users/joelneft/.codex/workspaces/default/"
    "AI for Teachers - Participant Workbook (Session 1) polished final.docx"
)


NAVY = "1F3B63"
TEAL = "2E6F73"
GOLD = "B8872B"
LIGHT = "F4F7FA"
LIGHT_ALT = "EAF1F4"
WHITE = "FFFFFF"
TEXT = RGBColor(34, 34, 34)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def format_run(run, size=None, bold=False, color=TEXT, italic=False):
    font = run.font
    font.name = "Aptos"
    font.size = Pt(size or 10.5)
    font.bold = bold
    font.italic = italic
    font.color.rgb = color


def make_paragraph(
    doc,
    text="",
    *,
    style=None,
    size=10.5,
    bold=False,
    color=TEXT,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    space_before=0,
    space_after=6,
    italic=False,
):
    paragraph = doc.add_paragraph(style=style)
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(space_before)
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.line_spacing = 1.05
    if text:
        run = paragraph.add_run(text)
        format_run(run, size=size, bold=bold, color=color, italic=italic)
    return paragraph


def make_heading(doc, text, level=1):
    if level == 1:
        return make_paragraph(
            doc,
            text,
            size=15,
            bold=True,
            color=RGBColor(31, 59, 99),
            space_before=8,
            space_after=6,
        )
    return make_paragraph(
        doc,
        text,
        size=12.5,
        bold=True,
        color=RGBColor(46, 111, 115),
        space_before=6,
        space_after=4,
    )


def section_banner(doc, part, title, subtitle):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade_cell(cell, NAVY)
    set_cell_margins(cell, top=130, bottom=130, start=140, end=140)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run(part)
    format_run(r1, size=11, bold=True, color=RGBColor(255, 255, 255))

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(title)
    format_run(r2, size=18, bold=True, color=RGBColor(255, 255, 255))

    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(subtitle)
    format_run(r3, size=10.5, color=RGBColor(255, 255, 255))

    make_paragraph(doc, "", space_after=2)


def info_box(doc, title, body, fill=LIGHT):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    set_cell_margins(cell, top=100, bottom=100, start=130, end=130)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    p_title = cell.paragraphs[0]
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_title = p_title.add_run(title)
    format_run(r_title, size=11.5, bold=True, color=RGBColor(31, 59, 99))

    p_body = cell.add_paragraph()
    p_body.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_body = p_body.add_run(body)
    format_run(r_body, size=10.5, color=TEXT)
    p_body.paragraph_format.space_after = Pt(0)
    make_paragraph(doc, "", space_after=3)


def two_col_table(doc, headers, rows, widths=None, fill_header=NAVY):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        cell = header_cells[i]
        cell.text = ""
        shade_cell(cell, fill_header)
        set_cell_margins(cell, top=100, bottom=100, start=120, end=120)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        r = p.add_run(header)
        format_run(r, size=10.5, bold=True, color=RGBColor(255, 255, 255))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if widths:
            cell.width = widths[i]
    set_repeat_table_header(table.rows[0])

    for row_idx, values in enumerate(rows):
        row_cells = table.add_row().cells
        for i, value in enumerate(values):
            cell = row_cells[i]
            cell.text = ""
            shade_cell(cell, WHITE if row_idx % 2 == 0 else LIGHT)
            set_cell_margins(cell, top=90, bottom=90, start=120, end=120)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for line_idx, line in enumerate(str(value).split("\n")):
                if line_idx == 0:
                    run = p.add_run(line)
                else:
                    p = cell.add_paragraph()
                    run = p.add_run(line)
                format_run(run, size=10.2, color=TEXT)
                p.paragraph_format.space_after = Pt(0)
            if widths:
                cell.width = widths[i]
    make_paragraph(doc, "", space_after=3)
    return table


def note_box(doc, title, lines=4, fill=WHITE):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    set_cell_margins(cell, top=90, bottom=90, start=130, end=130)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(title)
    format_run(r, size=10.7, bold=True, color=RGBColor(31, 59, 99))
    for _ in range(lines):
        line_p = cell.add_paragraph()
        line_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        line_p.paragraph_format.space_after = Pt(0)
        line_run = line_p.add_run("_" * 86)
        format_run(line_run, size=10.2, color=RGBColor(110, 110, 110))
    make_paragraph(doc, "", space_after=3)


def prompt_notes_table(doc, prompts):
    rows = [[prompt, ""] for prompt in prompts]
    return two_col_table(
        doc,
        ["Discussion Prompt", "Notes"],
        rows,
        fill_header=TEAL,
    )


def checklist(doc, items):
    for item in items:
        make_paragraph(
            doc,
            f"[ ] {item}",
            size=10.5,
            space_after=2,
        )
    make_paragraph(doc, "", space_after=4)


def add_cover(doc):
    make_paragraph(doc, "", space_after=28)
    make_paragraph(
        doc,
        "PROFESSIONAL LEARNING",
        size=11,
        bold=True,
        color=RGBColor(184, 135, 43),
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=10,
    )
    make_paragraph(
        doc,
        "AI for Accessible Instruction",
        size=24,
        bold=True,
        color=RGBColor(31, 59, 99),
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=6,
    )
    make_paragraph(
        doc,
        "Participant Workbook",
        size=16,
        bold=True,
        color=RGBColor(46, 111, 115),
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=4,
    )
    make_paragraph(
        doc,
        "Session 1: Access Without Lowering Demand",
        size=13,
        bold=False,
        color=TEXT,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=12,
    )

    info_box(
        doc,
        "Use this workbook to do four things well.",
        (
            "Identify one real access barrier, classify supports through the OSAMR lens, "
            "build one verified classroom-ready scaffold, and leave with a school-year launch plan."
        ),
        fill=LIGHT_ALT,
    )

    make_paragraph(
        doc,
        "Baltimore City Schools | EdTech 2026 | Facilitator: Joel Neft",
        size=10.5,
        color=RGBColor(90, 90, 90),
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=14,
        space_after=0,
    )
    doc.add_page_break()


def add_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.text = ""
    run = footer_para.add_run(
        "AI for Accessible Instruction | Session 1 Participant Workbook | Baltimore City Schools"
    )
    format_run(run, size=8.5, color=RGBColor(110, 110, 110))


def configure_doc():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = TEXT
    add_footer(doc)
    return doc


def build_document():
    doc = configure_doc()
    add_cover(doc)

    make_heading(doc, "The Challenge We're Solving")
    info_box(
        doc,
        "The problem",
        (
            "Teachers often struggle meeting the diverse needs of all students simultaneously. "
            "Without efficient tools, creating multiple access points for a single topic can be time-consuming, "
            'which can lead to a "one-size-fits-all" approach that leaves some students behind.'
        ),
    )
    info_box(
        doc,
        "Today's move",
        (
            "Use OSAMR as a checkpoint to distinguish stronger access from erosion of rigor. "
            "Remove one barrier while preserving the learning target, the cognitive verb, and the task demand."
        ),
        fill=LIGHT_ALT,
    )

    make_heading(doc, "What You'll Leave With Today")
    two_col_table(
        doc,
        ["#", "Takeaway"],
        [
            ["1", "A school-year ready, barrier-targeted support aligned to an upcoming lesson"],
            ["2", "A reusable build recipe: Input -> Output -> Verification"],
            ["3", "An OSAMR quick check explaining how rigor was protected"],
        ],
        fill_header=TEAL,
    )

    make_heading(doc, "Non-Negotiable Protocol")
    checklist(
        doc,
        [
            "Thinking preserved.",
            "Barrier removed.",
            "Learning target unchanged.",
        ],
    )
    info_box(
        doc,
        "Session filter",
        "Every activity, prompt, and tool choice in today's session should clear all three checks.",
        fill=LIGHT_ALT,
    )

    make_heading(doc, "Today's Agenda (2.5 Hours)")
    two_col_table(
        doc,
        ["Time", "Session Block", "Purpose"],
        [
            ["0-10 min", "Part 1 | Name Your Barrier", "Anchor the session in one real access problem."],
            ["10-35 min", "Part 2 | The OSAMR Framework", "Use cases to distinguish stronger access from erosion."],
            ["35-55 min", "Part 3 | Fix-It Design Challenge", "Redesign overreliance examples so the thinking stays."],
            ["55-125 min", "Part 4 | Build Sprint", "Choose one tool, build one support, and verify it."],
            ["125-150 min", "Part 5 | Commit & Share", "Showcase the support, reflect, and plan for school-year launch."],
        ],
    )
    doc.add_page_break()

    section_banner(doc, "PART 1", "Name Your Barrier", "10 minutes | Identify one barrier")
    make_heading(doc, "Welcome and Frame the Work", level=2)
    info_box(
        doc,
        "Today's goal",
        (
            "Leave with one ready-to-use support that removes a barrier while keeping the thinking intact. "
            "Think of one student who struggled with a recent task not because they could not think, "
            "but because something got in the way."
        ),
    )
    info_box(
        doc,
        "Quick share",
        (
            "Name the barrier. Then name the thinking students still need to do. "
            "Common barriers include reading level, unclear directions, unfamiliar vocabulary, unfamiliar format, or language demands."
        ),
        fill=LIGHT_ALT,
    )

    make_heading(doc, "Common Barriers", level=2)
    two_col_table(
        doc,
        ["Barrier", "What it often looks like"],
        [
            ["Reading level is too high", "Students cannot access the text well enough to begin the task independently."],
            ["Directions are unclear", "Students are unsure what to do, where to start, or what success should look like."],
            ["Vocabulary is unfamiliar", "Key terms block comprehension of the task or text."],
            ["Language demands are too heavy", "Students can think about the content, but need support expressing that thinking."],
            ["Task format is unfamiliar", "Students have not yet learned how to navigate the response structure or format."],
        ],
        fill_header=TEAL,
    )

    make_heading(doc, "Barrier Snapshot", level=2)
    two_col_table(
        doc,
        ["Timing", "Mode", "What to do"],
        [
            ["1 min", "Individual", "Silently identify one specific barrier in an upcoming lesson."],
            ["2 min", "Partner talk", "Share the barrier and explain what students still need to think about or do."],
            ["2 min", "Whole group", "Listen for patterns as common barriers are surfaced across the room."],
        ],
    )

    note_box(doc, "The one barrier I am addressing today", lines=4)
    note_box(doc, "The learning target I am protecting", lines=3, fill=LIGHT)
    note_box(doc, "The cognitive verb / hard thinking I still want students to do", lines=3)
    doc.add_page_break()

    section_banner(doc, "PART 2", "The OSAMR Framework", "25 minutes | Mini-lesson + case study analysis")
    make_heading(doc, "OSAMR Mini-Lesson", level=2)
    info_box(
        doc,
        "The big question",
        "Does AI increase access, or does it remove the productive struggle students still need?",
    )
    two_col_table(
        doc,
        ["Lens", "Question to ask"],
        [
            ["Barrier", "What barrier did the tool actually remove?"],
            ["Thinking", "What work still belongs to students?"],
            ["Verification", "What will the teacher still check?"],
        ],
        fill_header=TEAL,
    )

    make_heading(doc, "The OSAMR Levels", level=2)
    two_col_table(
        doc,
        ["Level", "What it means", "Quick example"],
        [
            ["O | Overreliance", "AI did the thinking for the student or removed the hard step.", "AI writes the essay or replaces the text with a summary."],
            ["S | Substitution", "The format changed, but the task stayed the same.", "A digital worksheet replaces a paper worksheet."],
            ["A | Augmentation", "A support feature improves access inside the same task.", "Text-to-speech supports reading while the task stays intact."],
            ["M | Modification", "The task design shifts meaningfully while the target still holds.", "Students use AI-supported planning tools but still complete the core reasoning themselves."],
            ["R | Redefinition", "Technology enables a rigorous task that was not otherwise possible.", "Students analyze patterns across multimodal sources in ways that were not previously practical."],
        ],
    )
    info_box(
        doc,
        "Key principle",
        "You do not need Redefinition. Right-sized support at any level beats flashy misalignment. The question is not how advanced the tool looks. The question is whether the support preserved the thinking.",
        fill=LIGHT_ALT,
    )

    make_heading(doc, "Case Study Analysis Structure", level=2)
    two_col_table(
        doc,
        ["Step", "Time", "Action"],
        [
            ["Model one scenario", "5 min", "Analyze one example together before groups begin."],
            ["Annotate in groups", "7 min", "Mark barrier removed, thinking preserved, and any red flags."],
            ["Classify + prepare", "3 min", "Agree on an OSAMR classification and one piece of evidence to share."],
        ],
        fill_header=TEAL,
    )

    make_heading(doc, "Case Study 1: Translated Directions", level=2)
    info_box(
        doc,
        "Grade 6 science | Barrier: language access | AI move: Spanish directions",
        (
            "Ms. Rodriguez teaches 6th grade science during a water cycle lesson. "
            "She uses AI to translate the lesson directions into Spanish for her MLL students. "
            "Students still diagram evaporation, condensation, and precipitation using evidence from the text. "
            "The tool changes the access point, not the scientific thinking."
        ),
    )
    prompt_notes_table(
        doc,
        [
            "What barrier was removed?",
            "What thinking still belongs to students?",
            "Which OSAMR level fits best and why?",
            "What evidence shows rigor stayed intact?",
        ],
    )

    make_heading(doc, "Case Study 2: AI Summary", level=2)
    info_box(
        doc,
        "Grade 9 ELA | Barrier: complex text access | AI move: summary replaces text",
        (
            "Mr. Johnson teaches 9th grade English. Students are reading a complex chapter from To Kill a Mockingbird. "
            "He asks AI to summarize the chapter in simpler language and gives students the summary instead of the original text. "
            "Students answer comprehension questions using the AI summary."
        ),
        fill=LIGHT_ALT,
    )
    prompt_notes_table(
        doc,
        [
            "What barrier was the teacher trying to remove?",
            "What is the problem with this move?",
            "Which OSAMR level fits best and why?",
            "How could the support be redesigned so the original thinking stays?",
        ],
    )

    make_heading(doc, "Whole-Group Share-Out", level=2)
    two_col_table(
        doc,
        ["Share-out lens", "Listen for red flags"],
        [
            [
                "Scenario summary\nOSAMR classification + evidence\nOne key insight or debate",
                "AI generates answers.\nA thinking step disappears.\nThe original task gets oversimplified.\nTeacher visibility into reasoning drops.",
            ]
        ],
        fill_header=TEAL,
    )
    note_box(doc, "What red flags for overreliance kept showing up across scenarios?", lines=4)
    doc.add_page_break()

    section_banner(doc, "PART 3", "Fix-It Design Challenge", "20 minutes | Redesign overreliance examples")
    make_heading(doc, "Fix-It Design Challenge", level=2)
    info_box(
        doc,
        "Redesign task",
        (
            "Keep the same barrier the original tool was trying to address. "
            "Redesign the support so the barrier is removed while the thinking stays with students."
        ),
    )
    make_heading(doc, "Keep These Intact", level=2)
    checklist(
        doc,
        [
            "Keep the standard intact.",
            "Keep the success criteria intact.",
            "Keep the cognitive verb intact.",
        ],
    )

    make_heading(doc, "Redesign Template", level=2)
    two_col_table(
        doc,
        ["Prompt", "Your notes"],
        [
            ["1. Original problem\nWhat barrier existed?", ""],
            ["2. Overreliance issue\nWhy did the original AI move fail?", ""],
            ["3. Redesign solution\nHow does the new support maintain high expectations?", ""],
            ["4. OSAMR classification\nWhat level fits the redesigned support?", ""],
        ],
        fill_header=TEAL,
    )

    make_heading(doc, "Gallery Walk", level=2)
    two_col_table(
        doc,
        ["Rapid gallery walk", "Feedback prompts"],
        [
            [
                "30-45 seconds per redesign\nName the barrier it addressed.\nName the move that preserved rigor.",
                "What barrier was this solving?\nWhat makes this redesign work?\nWhat move could I use in my classroom?",
            ]
        ],
    )
    make_heading(doc, "Top 5 High-Impact Moves", level=2)
    checklist(
        doc,
        [
            "Chunk multi-step directions with numbered visuals.",
            "Provide sentence frames for constructed response without giving answers.",
            "Preview vocabulary with visuals, then require students to apply terms in context.",
            "Translate directions without simplifying the task itself.",
            "Generate feedback prompts that push thinking instead of correcting answers.",
        ],
    )
    note_box(doc, "Strongest redesign move I want to borrow", lines=4, fill=LIGHT)
    doc.add_page_break()

    section_banner(doc, "PART 4", "Build Sprint", "70 minutes | Create your classroom-ready support")
    make_heading(doc, "Choose Your Build Tool", level=2)
    two_col_table(
        doc,
        ["Tool", "Best use"],
        [
            ["Gemini", "Summarize or outline text, simplify wording while keeping concepts, and generate sentence frames."],
            ["Copilot", "Draft similar supports inside Microsoft workflows and build prompts, outlines, or frames."],
            ["NotebookLM", "Organize source material, generate study guides, and create discussion questions from existing content."],
        ],
        fill_header=TEAL,
    )
    info_box(
        doc,
        "Tool choice guidance",
        "Use the approved tool you can coach confidently. Tool choice matters less than alignment and verification.",
        fill=LIGHT_ALT,
    )

    make_heading(doc, "Build Readiness Checklist", level=2)
    checklist(
        doc,
        [
            "What barrier will I address? Use the barrier you named earlier.",
            "What task or content will I input to AI? Directions, passage, prompt, or assignment.",
            "What do I expect AI to output? Chunked directions, translated text, sentence frames, or another support.",
            "How will I verify it worked? Barrier removed, thinking preserved, and still aligned to the target.",
            "Technology check: Sign in to your chosen tool before the sprint starts.",
        ],
    )

    make_heading(doc, "Tool-Based Breakout Rooms", level=2)
    two_col_table(
        doc,
        ["Why tool-specific rooms?", "Room setup"],
        [
            [
                "Participants spend less time learning tools they are not using.\nCoaching is more targeted when everyone shares the same platform.\nPeer support improves when the room uses the same tool.",
                "Gemini Room\nCopilot Room\nNotebookLM Room\nJoin the room that matches the tool you are using.",
            ]
        ],
        fill_header=TEAL,
    )

    make_heading(doc, "Build Sprint 1: Create", level=2)
    info_box(
        doc,
        "Build focus",
        (
            "Make the support, then check it in real time. The goal is one verified support, not lots of rough drafts. "
            "Name the OSAMR level you are aiming for and use that lens to spot drift toward overreliance."
        ),
    )
    checklist(
        doc,
        [
            "Does this remove the barrier?",
            "Are students still doing the hard thinking?",
            "Is it aligned to the learning target?",
            "What will I verify before classroom use?",
        ],
    )

    make_heading(doc, "Prompt Templates", level=2)
    two_col_table(
        doc,
        ["Template 1 | Simplify Directions", "Template 2 | Create Sentence Starters"],
        [
            [
                "Rewrite these directions for [grade level] students. Break them into numbered steps. Use simple sentences. Keep the thinking work the same. Do NOT simplify the task, only the wording.\n\n[Paste your directions here]",
                "Create three sentence starters to help students [analyze / compare / justify]. Guide their thinking process, but do NOT give answers. Students must complete the reasoning. Learning target: [paste target].",
            ],
            [
                "Example barrier: directions unclear\nExample output: numbered steps in simpler language",
                "Example barrier: language demands\nExample output: sentence frames with space for student reasoning",
            ],
        ],
        fill_header=GOLD,
    )
    two_col_table(
        doc,
        ["Template 3 | Chunk Dense Text", "Template 4 | Translate Directions for MLL Support"],
        [
            [
                "Break this passage into three chunks. Add a focus question before each chunk. Do NOT simplify the text. Keep the original wording.\n\n[Paste passage]",
                "Translate these directions into [language]. Keep academic vocabulary in English with a translation in parentheses.\n\n[Paste directions]",
            ],
            [
                "Best use: dense reading without replacing the source text",
                "Best use: multilingual access while protecting task rigor",
            ],
        ],
        fill_header=GOLD,
    )

    make_heading(doc, "Verification Checkpoint", level=2)
    info_box(
        doc,
        "Non-negotiable check",
        "Every support is checked before participants leave the room. Barrier removed. Thinking preserved. Alignment confirmed.",
        fill=LIGHT_ALT,
    )
    two_col_table(
        doc,
        ["Status", "Meaning"],
        [
            ["READY TO USE", "Barrier removed. Thinking preserved. Alignment confirmed. Proceed to documentation."],
            ["REFINE & STRENGTHEN", "Close, but one adjustment is still needed. Make the change now."],
            ["REWORK FOR ALIGNMENT", "The move may oversimplify the task or drift from the target. Brief coaching plus revision."],
        ],
        fill_header=GOLD,
    )
    info_box(
        doc,
        "Shared tracker",
        "The tracker should make it visible who is ready, refining, or reworking so participants know where to focus their energy.",
    )
    make_heading(doc, "The 3 Verification Questions", level=2)
    two_col_table(
        doc,
        ["Question", "What to test"],
        [
            ["1. Did it remove the barrier?", "Can the student who was previously stuck access the task more effectively now?"],
            ["2. Do students still do the thinking?", "Does the learning target verb still require student work?"],
            ["3. Does it match the standard?", "Is the cognitive demand preserved?"],
        ],
        fill_header=TEAL,
    )
    info_box(
        doc,
        "Real-time quality check",
        "Test the support with one sample before finalizing it. Look for mistakes, hallucinations, weak translations, or hidden simplification.",
    )

    make_heading(doc, "Build Sprint 2: Refine & Document", level=2)
    two_col_table(
        doc,
        ["Documentation prompt", "Your notes"],
        [
            ["1. Input\nWhat material did I give the tool?", ""],
            ["2. Process\nWhat did the AI generate?", ""],
            ["3. Verification\nHow did I check that access improved and rigor stayed intact?", ""],
            ["4. OSAMR Lens\nWhat level fits and how did I avoid overreliance?", ""],
        ],
        fill_header=TEAL,
    )
    info_box(
        doc,
        "Why documentation matters",
        "Documentation makes the workflow reusable, shareable, and easier to improve next time.",
        fill=LIGHT_ALT,
    )
    note_box(doc, "My classroom-ready support in one sentence", lines=4, fill=LIGHT)
    doc.add_page_break()

    section_banner(doc, "PART 5", "Commit & Share", "25 minutes | Showcase + accountability")
    make_heading(doc, "Peer Showcase", level=2)
    two_col_table(
        doc,
        ["As each support is shared", "Capture one move to borrow"],
        [
            [
                "Name the barrier it solves.\nNotice the scaffold move.\nLook for where rigor is protected.\nConsider what evidence you would want from student work.",
                "Type one strength you notice.\nAdd one question or affirmation.\nName one move you could use in your own classroom.\nLook for ideas that transfer across grades or subjects.",
            ]
        ],
        fill_header=TEAL,
    )
    info_box(
        doc,
        "Showcase lens",
        "Use the showcase to notice what strong alignment looks like in practice.",
    )
    note_box(doc, "Peer showcase notes", lines=5)

    make_heading(doc, "OSAMR Self-Reflection", level=2)
    two_col_table(
        doc,
        ["Prompt", "Notes"],
        [
            ["1. What OSAMR level does your support represent?", ""],
            ["2. Why did you classify it there? Use one or two pieces of evidence.", ""],
            ["3. How did you ensure you avoided overreliance?", ""],
        ],
        fill_header=TEAL,
    )

    make_heading(doc, "Implementation Commitment Card", level=2)
    note_box(doc, "1. Lesson name + date for school-year launch", lines=3)
    note_box(doc, "2. One thing I will check in student work to know this helped", lines=3, fill=LIGHT)
    note_box(doc, "3. Accountability partner name + contact", lines=3)

    make_heading(doc, "Find an Accountability Partner", level=2)
    two_col_table(
        doc,
        ["Partner selection", "Commit to a check-in"],
        [
            [
                "Choose someone teaching a similar grade band, subject, or learning target.\nExchange an email or phone number before the session ends.",
                "After you both try the support, ask:\nDid it work?\nWhat did you see in student work?\nWhat would you change next time?",
            ]
        ],
        fill_header=TEAL,
    )

    make_heading(doc, "Important Reminders", level=2)
    checklist(
        doc,
        [
            "Do not include student names or protected information.",
            "Use approved tools and district guidance.",
            "Start with one barrier and one task.",
            "Test the support with one student sample first.",
            "Check for mistakes, hallucinations, or translation errors.",
            "Verify alignment to the standard and HQIM.",
            "You are the expert. AI is your tool.",
        ],
    )

    make_heading(doc, "Resources & Next Steps", level=2)
    info_box(
        doc,
        "Closing message",
        "Everything from today should leave participants with one verified support they can use immediately.",
        fill=LIGHT_ALT,
    )
    two_col_table(
        doc,
        ["Session assets", "How to use them"],
        [
            ["Prompt templates", "Use these as starting points when you need to clarify directions, create sentence starters, chunk text, or translate directions while protecting the task."],
            ["Case studies", "Return to the examples from Part 2 when you want a quick model of stronger access versus overreliance."],
            ["OSAMR quick-check", "Use it after generating a support to confirm that the barrier was reduced and the thinking stayed with students."],
            ["Build recipe notes", "Reuse your documented process so the next support is faster, clearer, and easier to improve."],
        ],
        fill_header=TEAL,
    )
    two_col_table(
        doc,
        ["Example tools", "Best use"],
        [
            ["Gemini", "Draft supports, clarify directions, and generate sentence frames or outlines."],
            ["Copilot", "Build similar supports inside Microsoft workflows."],
            ["NotebookLM", "Organize source material and generate study guides or discussion questions."],
        ],
        fill_header=TEAL,
    )
    two_col_table(
        doc,
        ["School-year launch", "Action"],
        [
            ["1", "Try one support in an opening lesson."],
            ["2", "Collect one student sample or observation note."],
            ["3", "Check the support against the three verification questions."],
            ["4", "Debrief with your partner and note one revision for next time."],
        ],
        fill_header=TEAL,
    )

    doc.save(OUTPUT_PATH)


if __name__ == "__main__":
    build_document()
